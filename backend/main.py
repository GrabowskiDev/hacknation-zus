from __future__ import annotations

import base64
import json
import os
from datetime import date
from enum import Enum
from typing import Any, List, Optional
from unittest import result

import io
import zipfile
from docx import Document as DocxDocument
from docx.shared import Pt
from dotenv import load_dotenv
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fpdf import FPDF
from langchain_core.output_parsers import PydanticOutputParser, StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_google_genai import ChatGoogleGenerativeAI
from pydantic import BaseModel, Field

import re
from pypdf import PdfReader, PdfWriter
from pypdf.generic import NameObject, BooleanObject, TextStringObject, DictionaryObject, ArrayObject

from ocr import (
    extract_text_from_image,
    summarize_accident_facts_from_pdfs,
    build_filled_card_text_from_summary,
)


load_dotenv()  # load GOOGLE_API_KEY and friends from .env

ACTION_PLAN_GUIDE = """Źródło: \"ZUS Accident Notification Tool\" (wytyczne dotyczące finalnych kroków po złożeniu zawiadomienia o wypadku).

1. Najpilniejsze działania po wypadku
- Zgłoś zdarzenie do Zakładu Ubezpieczeń Społecznych (PUE ZUS lub w oddziale) i zarejestruj sprawę.
- Zapewnij pełną dokumentację medyczną: zaświadczenia OL-9/N-9, karty wypisowe ze szpitala, zalecenia lekarza.
- Zbierz zeznania świadków oraz zdjęcia/notatki potwierdzające okoliczności wypadku.

2. Wymagane formularze i dokumenty według profilu zgłaszającego
Przedsiębiorca / osoba samozatrudniona:
- Zawiadomienie o wypadku (generowane w systemie) oraz druk ZUS Z-3b.
- Szczegółowa karta wypadku opisująca przebieg i przyczyny.
- Zaświadczenia medyczne OL-9 lub N-9 oraz dowód niezdolności do pracy.
- Potwierdzenie opłacania składek i podstawowe identyfikatory działalności (NIP, REGON, PKD).

Pracownik / zleceniobiorca:
- Niezwłoczne zgłoszenie wypadku pracodawcy i udział w sporządzeniu protokołu powypadkowego lub karty wypadku.
- Pracodawca wystawia druk ZUS Z-3 (pracownik) lub ZUS Z-3a (zleceniobiorca). Poszkodowany składa ZUS Z-15 przy wniosku o zasiłek.
- Jeśli interweniowały służby, dołącz notatki policji lub pogotowia.

Wypadek w drodze do lub z pracy:
- Wypełnij z pracodawcą lub przedstawicielem ZUS kartę wypadku w drodze.
- Dostarcz dowody potwierdzające trasę i cel podróży (bilety, logi GPS, zeznania świadków).

3. Przebieg po złożeniu zawiadomienia
Krok 1: Dokończ i podpisz dokumentację (zawiadomienie, wyjaśnienia, lista świadków, karta wypadku). Sprawdź spójność godzin pracy i momentu wypadku.
Krok 2: Złóż komplet dokumentów w ZUS i zachowaj potwierdzenie nadania. Kopie przechowuj w aktach sprawy.
Krok 3: Monitoruj status w PUE ZUS i reaguj na prośby o uzupełnienia. Dostarczaj dodatkowe zaświadczenia medyczne bez zwłoki.
Krok 4: Po decyzji ZUS sprawdź należne świadczenia (zasiłek chorobowy, świadczenie rehabilitacyjne, odszkodowanie). W razie potrzeby złóż odwołanie w ustawowym terminie.

4. Dodatkowe wskazówki z PDF
- Przy każdym kroku wypisz wymagane załączniki, by uniknąć ponownych wizyt w ZUS.
- Dbaj o zgodność treści między opisem zdarzenia, zeznaniami świadków i dokumentacją medyczną.
- Przypomnij poszkodowanemu o konieczności przechowywania dowodów opłaconych składek oraz instruktażu BHP otrzymanego przed wypadkiem.
- Przygotuj checklistę formularzy: zawiadomienie, ZUS Z-3/Z-3b/Z-3a, ZUS Z-15, karta wypadku, OL-9/N-9, wypisy medyczne, zeznania świadków, zdjęcia, potwierdzenie zabezpieczenia miejsca wypadku.
"""


class Mode(str, Enum):
    NOTIFICATION = "notification"
    EXPLANATION = "explanation"


class Witness(BaseModel):
    first_name: str = Field(..., description="Imię świadka")
    last_name: str = Field(..., description="Nazwisko świadka")
    address: Optional[str] = Field(None, description="Adres świadka")

class ReporterType(str, Enum):
    VICTIM = "victim"
    PROXY = "proxy"

class CaseState(BaseModel):
    reporter_type: Optional[ReporterType] = None # Kto zgłasza?
    proxy_document_attached: bool = False

    # Dane poszkodowanego
    pesel: Optional[str] = None
    first_name: Optional[str] = None
    last_name: Optional[str] = None
    date_of_birth: Optional[str] = None
    address_home: Optional[str] = None
    address_correspondence: Optional[str] = None

    # Dane działalności
    nip: Optional[str] = None
    regon: Optional[str] = None
    business_address: Optional[str] = None
    pkd: Optional[str] = None
    business_description: Optional[str] = None

    # Informacje o wypadku
    accident_date: Optional[str] = None
    accident_time: Optional[str] = None
    accident_place: Optional[str] = None
    planned_work_start: Optional[str] = None
    planned_work_end: Optional[str] = None
    injury_type: Optional[str] = None
    accident_description: Optional[str] = None
    first_aid_info: Optional[str] = None
    proceedings_info: Optional[str] = None
    equipment_info: Optional[str] = None

    # Świadkowie
    witnesses: List[Witness] = Field(default_factory=list)

    # Flagi definicyjne
    sudden: Optional[bool] = None
    external_cause: Optional[bool] = None
    injury_confirmed: Optional[bool] = None
    work_related: Optional[bool] = None


class MissingField(BaseModel):
    field: str
    reason: str


class ChatTurn(BaseModel):
    role: str  # "user" lub "assistant"
    content: str


class AssistantMessageRequest(BaseModel):
    case_id: str
    message: str
    mode: Mode
    conversation_history: List[ChatTurn] = Field(
        default_factory=list,
        description="Dotychczasowa historia rozmowy dla tej sprawy",
    )
    case_state: Optional[CaseState] = Field(
        default=None,
        description="Aktualny stan sprawy utrzymywany po stronie frontendu (opcjonalny)",
    )


class ActionStep(BaseModel):
    step_number: int = Field(..., description="Numer kolejny kroku")
    description: str = Field(..., description="Opis czynności do wykonania")
    required_documents: List[str] = Field(default_factory=list, description="Lista wymaganych dokumentów dla tego kroku")


class ActionPlan(BaseModel):
    actions: List[ActionStep] = Field(default_factory=list, description="Lista kroków do wykonania")


class AssistantMessageResponse(BaseModel):
    assistant_reply: str
    missing_fields: List[MissingField]
    case_state_preview: CaseState
    recommended_actions: Optional[List[ActionStep]] = None


class CaseDocument(BaseModel):
    """
    Pojedynczy dokument wejściowy do oceny sprawy (np. zawiadomienie, wyjaśnienia, karta informacyjna).
    """

    name: str = Field(..., description="Nazwa/tytuł dokumentu, np. 'zawiadomienie ZUS Z-3'")
    type: Optional[str] = Field(
        default=None,
        description="Rodzaj dokumentu, np. 'zawiadomienie', 'wyjaśnienia', 'dokumentacja medyczna'",
    )
    text: str = Field(..., description="Pełna treść dokumentu po odczytaniu (np. z OCR)")


class Discrepancy(BaseModel):
    """
    Rozbieżność pomiędzy dokumentami lub pomiędzy dokumentem a ustalonym stanem faktycznym.
    """

    description: str = Field(
        ...,
        description=(
            "Opis rozbieżności (np. różne daty wypadku, różne miejsce wypadku, niespójne dane świadka)"
        ),
    )
    fields_affected: List[str] = Field(
        default_factory=list,
        description="Nazwy pól, których dotyczy rozbieżność (np. 'accident_date', 'accident_place')",
    )
    documents_involved: List[str] = Field(
        default_factory=list,
        description="Nazwy dokumentów, w których występuje rozbieżność",
    )


class OpinionOutcome(str, Enum):
    """
    Wynik oceny zdarzenia.
    """

    ACCIDENT_CONFIRMED = "accident_confirmed"
    ACCIDENT_NOT_CONFIRMED = "accident_not_confirmed"
    INCONCLUSIVE = "inconclusive"


class CaseEvaluationResult(BaseModel):
    """
    Wynik złożonej oceny sprawy na podstawie całego zestawu dokumentów.
    """

    normalized_case_state: CaseState = Field(
        ...,
        description=(
            "Ujednolicony stan faktyczny sprawy wynikający z całości dokumentów (odpowiada polom CaseState)"
        ),
    )
    discrepancies: List[Discrepancy] = Field(
        default_factory=list,
        description="Lista rozbieżności wykrytych w dokumentach",
    )
    missing_fields: List[MissingField] = Field(
        default_factory=list,
        description="Pola, które nadal wymagają uzupełnienia (np. brakujące informacje w dokumentacji)",
    )
    missing_documents: List[str] = Field(
        default_factory=list,
        description="Informacje o dokumentach, które należy jeszcze uzupełnić/dostarczyć",
    )
    opinion: OpinionOutcome = Field(
        ...,
        description=(
            "Czy zdarzenie należy uznać za wypadek podczas prowadzenia pozarolniczej działalności gospodarczej"
        ),
    )
    opinion_explanation: str = Field(
        ...,
        description="Szczegółowe, logiczne uzasadnienie przyjętego rozstrzygnięcia",
    )
    accident_card_draft: str = Field(
        ...,
        description=(
            "Projekt karty wypadku w formie tekstowej (może być w formacie zbliżonym do formularza ZUS)"
        ),
    )


class CaseEvaluationRequest(BaseModel):
    case_id: str = Field(..., description="Identyfikator sprawy (dowolny identyfikator z frontendu)")
    documents: List[CaseDocument] = Field(
        default_factory=list,
        description="Komplet dokumentów tekstowych (np. zawiadomienia, wyjaśnienia, dokumentacja medyczna)",
    )


class CaseEvaluationResponse(BaseModel):
    case_id: str
    evaluation: CaseEvaluationResult


app = FastAPI(title="ZANT Backend", version="0.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # na potrzeby hackathonu puszczamy wszystko
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["X-ZANT-Actions"],
)


@app.get("/")
async def root() -> dict:
    return {"status": "ok", "service": "ZANT backend"}


def get_llm() -> Optional[Any]:
    """
    Prosty factory na LLM-a.
    Jeśli LangChain/Gemini nie są zainstalowane, zwracamy None,
    a pipeline zadziała w trybie fallback (bez LLM).
    """
    # Możesz sterować modelem przez ENV: GEMINI_MODEL=gemini-1.5-flash
    return ChatGoogleGenerativeAI(
        model=os.getenv("GEMINI_MODEL", "gemini-2.5-flash"),
        temperature=0,
    )


def simple_missing_fields(
    case_state: CaseState,
    mode: Mode,
    skipped_fields: Optional[List[str]] = None,
) -> List[MissingField]:
    """
    Bardzo prosty checker braków.
    Docelowo tu możesz zakodować wymagane pola dla różnych trybów.
    """
    required_common = ["accident_date", "accident_place", "accident_description"]
    required_zawiadomienie = ["injury_type"]
    required_wyjasnienia: List[str] = []

    required = set(required_common)
    if mode == Mode.NOTIFICATION:
        required.update(required_zawiadomienie)
    else:
        required.update(required_wyjasnienia)

    skipped = set(skipped_fields or [])
    result: List[MissingField] = []

    if case_state.reporter_type == ReporterType.PROXY and not case_state.proxy_document_attached:
        result.append(MissingField(
            field="proxy_document", 
            reason="W przypadku zgłoszenia przez pełnomocnika wymagane jest załączenie pełnomocnictwa."
        ))

    for field_name in required:
        if field_name in skipped:
            continue
        value = getattr(case_state, field_name, None)
        if value is None or (isinstance(value, str) and not value.strip()):
            result.append(MissingField(field=field_name, reason="brak wartości"))
        elif (
            isinstance(value, str)
            and len(value.strip()) < 20
            and field_name == "accident_description"
        ):
            result.append(
                MissingField(field=field_name, reason="opis jest zbyt ogólny")
            )

    return result


FIELD_LABELS: dict[str, str] = {
    # Dane poszkodowanego
    "first_name": "imię poszkodowanego",
    "last_name": "nazwisko poszkodowanego",
    "pesel": "PESEL poszkodowanego",
    "date_of_birth": "data urodzenia poszkodowanego",
    "address_home": "adres zamieszkania poszkodowanego",
    "address_correspondence": "adres do korespondencji",
    # Dane działalności
    "nip": "NIP działalności",
    "regon": "REGON działalności",
    "business_address": "adres prowadzenia działalności",
    "pkd": "kod PKD działalności",
    "business_description": "opis rodzaju prowadzonej działalności",
    # Informacje o wypadku
    "accident_date": "data wypadku",
    "accident_time": "godzina wypadku",
    "accident_place": "miejsce wypadku",
    "planned_work_start": "planowana godzina rozpoczęcia pracy",
    "planned_work_end": "planowana godzina zakończenia pracy",
    "injury_type": "rodzaj urazu",
    "accident_description": "opis okoliczności i przyczyn wypadku",
    "first_aid_info": "informacje o udzielonej pierwszej pomocy / szpitalu",
    "proceedings_info": "informacje o postępowaniu (policja, prokuratura itp.)",
    "equipment_info": "informacje o maszynach/urządzeniach, BHP, środkach ochrony",
}

# Kategorie pytań – grupujemy pola w logiczne sekcje.
CATEGORY_DEFS: list[tuple[str, str, list[str]]] = [
    (
        "person",
        "dane poszkodowanego",
        [
            "first_name",
            "last_name",
            "pesel",
            "date_of_birth",
            "address_home",
            "address_correspondence",
        ],
    ),
    (
        "business",
        "dane działalności",
        ["nip", "regon", "business_address", "pkd", "business_description"],
    ),
    (
        "accident",
        "informacje o wypadku",
        [
            "accident_date",
            "accident_time",
            "accident_place",
            "planned_work_start",
            "planned_work_end",
            "injury_type",
            "accident_description",
            "first_aid_info",
            "proceedings_info",
            "equipment_info",
        ],
    ),
]


def human_field_label(field_name: str) -> str:
    """
    Przyjazne nazwy pól do komunikatu dla użytkownika.
    """
    return FIELD_LABELS.get(field_name, field_name)


def find_category_for_field(field_name: str) -> Optional[tuple[str, str, list[str]]]:
    for cat in CATEGORY_DEFS:
        if field_name in cat[2]:
            return cat
    return None


def find_category_by_label(label: str) -> Optional[tuple[str, str, list[str]]]:
    label_norm = label.strip().lower()
    for cat in CATEGORY_DEFS:
        _, category_label, _ = cat
        if category_label.lower() == label_norm:
            return cat
    return None


def field_name_from_label(label: str) -> Optional[str]:
    for key, value in FIELD_LABELS.items():
        if value == label:
            return key
    return None


def message_looks_like_skip(text: str) -> bool:
    """
    Bardzo prosta heurystyka: czy użytkownik chce pominąć odpowiedź.
    """
    t = text.strip().lower()
    if not t:
        return False
    phrases = [
        "nie chcę podawać",
        "nie chce podawać",
        "nie chcę tego podawać",
        "nie chce tego podawać",
        "nie chcę podać",
        "nie chce podać",
        "nie podam",
        "wolę nie podawać",
        "wole nie podawac",
        "wolę nie mówić",
        "wolę nie mowic",
        "nie chcę mówić",
        "nie chce mowic",
        "pomiń",
        "pomin",
        "pomijamy",
        "pomiń kategorię",
        "idź dalej z kategorią",
        "następna kategoria"
    ]
    return any(p in t for p in phrases)


def message_asks_next_category(text: str) -> bool:
    """
    Czy użytkownik prosi, żeby przejść do kolejnej kategorii pytań.
    """
    t = text.strip().lower()
    if not t:
        return False
    phrases = [
        "następna kategoria",
        "nastepna kategoria",
        "kolejna kategoria",
        "idź dalej z kategorią",
        "idz dalej z kategoria",
        "pomiń kategorię",
        "pomijam tę kategorię",
        "chcę pominąć kategorię",
        "chcialbym pominac te kategorie",
        "chciałbym pominąć tę kategorię",
        "przejdź dalej bez tej kategorii",
        "przejdz dalej bez tej kategorii",
    ]
    return any(p.lower() in t.lower() for p in phrases)


def extract_category_label_from_text(text: str) -> Optional[str]:
    """
    Próbujemy wyciągnąć nazwę kategorii z tekstu pytania asystenta.
    Szukamy wzorców typu 'kategorii: X.' lub 'kategoria: X.'.
    """
    lowered = text.lower()
    for marker in ["kategorii:", "kategoria:"]:
        idx = lowered.find(marker)
        if idx != -1:
            start = idx + len(marker)
            # Wytnij do końca zdania / linii
            rest = text[start:].strip()
            for sep in [".", "\n"]:
                sep_idx = rest.find(sep)
                if sep_idx != -1:
                    rest = rest[:sep_idx]
                    break
            return rest.strip()
    return None


def detect_skip_with_llm(question_label: str, answer: str) -> bool:
    """
    Używa LLM do wykrycia, czy użytkownik odmawia podania danej informacji.
    Jeśli LLM nie jest dostępny, spada do prostej heurystyki.
    """
    llm = get_llm()
    if llm is None or ChatPromptTemplate is None:
        return message_looks_like_skip(answer)

    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                (
                    "Jesteś klasyfikatorem odpowiedzi użytkownika.\n"
                    "Masz ustalić, czy użytkownik NIE poda tej informacji w dalszym procesie, "
                    "bo albo wyraźnie odmawia, albo mówi, że taka informacja go nie dotyczy "
                    "(np. nie prowadzi działalności, nie ma NIP-u, nie pamięta daty).\n"
                    "Jeśli użytkownik odmawia podania informacji lub mówi, że jej nie posiada / "
                    "go nie dotyczy, odpowiedz dokładnie: YES.\n"
                    "Jeśli użytkownik próbuje odpowiedzieć (nawet nieprecyzyjnie, ale wygląda na to, "
                    "że chce udzielić informacji), odpowiedz dokładnie: NO.\n"
                    "Nie tłumacz się, nie dodawaj komentarzy, tylko jedno słowo: YES albo NO."
                ),
            ),
            (
                "human",
                (
                    "Pytanie asystenta (o jaką daną chodzi): {label}\n"
                    "Odpowiedź użytkownika: {answer}"
                ),
            ),
        ]
    )

    chain = prompt | llm | StrOutputParser()

    result = chain.invoke({"label": question_label, "answer": answer}).strip().upper()
    return result.startswith("YES")


def infer_skipped_fields_from_history(history: List[ChatTurn]) -> List[str]:
    """
    Przechodzi po historii:
    - gdy asystent pyta o kategorię lub konkretne pole,
    - a kolejna odpowiedź użytkownika wygląda jak odmowa / „to mnie nie dotyczy”
      lub prośba o przejście dalej,
    - oznaczamy odpowiednie pola jako "skipped".
    """
    skipped: set[str] = set()
    for i in range(len(history) - 1):
        turn = history[i]
        next_turn = history[i + 1]
        if turn.role != "assistant" or next_turn.role != "user":
            continue
        text = turn.content
        answer = next_turn.content

        # Najpierw spróbujmy zinterpretować to jako pytanie o kategorię.
        cat_label = extract_category_label_from_text(text)
        if cat_label:
            cat = find_category_by_label(cat_label)
            if cat and (
                message_asks_next_category(answer)
                or detect_skip_with_llm(cat_label, answer)
            ):
                _, _, category_fields = cat
                skipped.update(category_fields)
            continue

        # Jeśli to nie wygląda jak pytanie o kategorię, spróbujmy potraktować je
        # jak pytanie o konkretne pole (np. pole z etykietą).
        if ":" in text:
            label = (
                text.split(":")[-2 if text.count(":") > 1 else 0]
                .splitlines()[-1]
                .strip()
            )
            field_name = field_name_from_label(label)
            if field_name and detect_skip_with_llm(label, answer):
                skipped.add(field_name)
    return list(skipped)


def extract_case_state_with_llm(
    previous_state: CaseState,
    message: str,
    mode: Mode,
    today: str,
    conversation_history: List[ChatTurn],
) -> CaseState:
    """
    Wykorzystuje LangChain + LLM (Gemini) do uzupełnienia CaseState na podstawie wiadomości.
    Jeśli Gemini nie jest dostępny, działa w trybie fallback.
    """
    llm = get_llm()
    if llm is None or ChatPromptTemplate is None or PydanticOutputParser is None:
        # Fallback: tylko podmień opis wypadku na podstawie wiadomości
        return previous_state.model_copy(
            update={
                "accident_description": message.strip()
                or previous_state.accident_description
            }
        )


    parser = PydanticOutputParser(pydantic_object=CaseState)

    history_text = "\n".join(
        f"{turn.role}: {turn.content}" for turn in conversation_history[-10:]
    )

    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                (
                    "Jesteś asystentem pomagającym wypełnić dane o wypadku "
                    "dla ZUS. Uzupełniasz strukturalne pola w JSON-ie "
                    "na podstawie rozmowy z użytkownikiem.\n\n"
                    "Dzisiejsza data (czas serwera backendu): {today}\n\n"
                    "Zawsze zwracaj pełny JSON zgodny ze schematem CaseState.\n"
                    "Jeśli czegoś nie wiesz, pozostaw dane pola bez zmian."
                ),
            ),
            (
                "human",
                (
                    "Aktualny stan danych (CaseState) w formacie JSON:\n"
                    "{current_state}\n\n"
                    "Historia rozmowy (ostatnie wiadomości):\n"
                    "{history}\n\n"
                    "Tryb: {mode}\n\n"
                    "Nowa wiadomość użytkownika:\n"
                    "{message}\n\n"
                    "Twoje zadanie:\n"
                    "- wywnioskuj i uzupełnij tylko te pola, które można "
                    "jednoznacznie określić na podstawie rozmowy,\n"
                    "- pole reporter_type może mieć tylko wartości 'victim' (poszkodowany) lub 'proxy' (pełnomocnik),\n"
                    "- adresy address_home oraz address_correspondence zapisuj w formacie 'ulica nr/lokal, 00-000 Miasto' (np. 'ul. Przykładowa 12/4, 12-343 Warszawa'),\n"
                    "- pozostałe pola pozostaw bez zmian (przepisz ich "
                    "dotychczasową wartość),\n"
                    "- odpowiedz wyłącznie JSON-em pasującym do schematu CaseState."
                ),
            ),
        ]
    )

    chain = prompt | llm | parser

    try:
        updated_state: CaseState = chain.invoke(
            {
                "current_state": previous_state.model_dump(),
                "mode": mode.value,
                "message": message,
                "today": today,
                "history": history_text,
            }
        )
        return updated_state
    except Exception as e:
        # Fallback: tylko podmień opis wypadku na podstawie wiadomości
        print(f"BŁĄD LLM: {e}") 

        return previous_state.model_copy(
            update={
                "accident_description": message.strip()
                or previous_state.accident_description
            }
        )
def generate_post_accident_actions(case_state: CaseState) -> List[ActionStep]:
    """
    Generuje spersonalizowaną listę kroków i dokumentów na podstawie zebranych danych.
    """
    llm = get_llm()
    if llm is None:
        return []

    guide_excerpt = ACTION_PLAN_GUIDE

    # Używamy PydanticOutputParser z wrapperem ActionPlan, aby uzyskać poprawną strukturę listy.
    parser = PydanticOutputParser(pydantic_object=ActionPlan)

    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                (
                    "Jesteś ekspertem ZUS i wirtualnym doradcą w systemie ZANT (ZUS Accident Notification Tool). "
                    "Twoim zadaniem jest wygenerowanie spersonalizowanej listy czynności (Action Plan) dla użytkownika, "
                    "który zgłasza wypadek.\n\n"
                    "ŹRÓDŁO REFERENCYJNE (fragment PDF 'ZUS Accident Notification Tool'):\n"
                    "{guide_excerpt}\n\n"
                    "ANALIZA SYTUACJI (na podstawie CaseState):\n"
                    "1. Sprawdź, czy poszkodowany jest Przedsiębiorcą (wypełnione NIP, REGON, business_description) czy Pracownikiem.\n"
                    "2. Sprawdź rodzaj wypadku (przy pracy, w drodze, inny).\n\n"
                    "ZASADY GENEROWANIA KROKÓW (zgodnie z wytycznymi ZUS):\n"
                    "- Jeśli to PRZEDSIĘBIORCA (Wypadek przy pracy):\n"
                    "  1. Zgłoszenie wypadku w placówce ZUS (pisemnie lub przez PUE).\n"
                    "  2. Złożenie wniosku o ustalenie okoliczności wypadku (Karta Wypadku).\n"
                    "  3. Skompletowanie dokumentacji medycznej (zaświadczenie OL-9 lub N-9).\n"
                    "  4. Wypełnienie druku ZUS Z-3b (zaświadczenie płatnika składek).\n"
                    "- Jeśli to PRACOWNIK:\n"
                    "  1. Zgłoszenie wypadku pracodawcy (niezwłocznie).\n"
                    "  2. Udział w sporządzeniu Protokołu Powypadkowego (przez zespół powypadkowy).\n"
                    "  3. Przekazanie pracodawcy druku ZUS Z-15 (wniosek o zasiłek) i ZUS Z-3 (wypełnia pracodawca).\n"
                    "- Jeśli WYPADEK W DRODZE DO PRACY:\n"
                    "  1. Zgłoszenie pracodawcy/zleceniodawcy.\n"
                    "  2. Karta wypadku w drodze do pracy.\n\n"
                    "WYMAGANIA DLA ODPOWIEDZI:\n"
                    "- Zwróć obiekt JSON zgodny ze schematem ActionPlan (zawierający listę 'actions').\n"
                    "- Pola kroku: 'step_number' (int), 'description' (str), 'required_documents' (list[str]).\n"
                    "- Bądź precyzyjny i pomocny. Używaj oficjalnych nazw druków ZUS.\n\n"
                    "{format_instructions}"
                ),
            ),
            (
                "human",
                (
                    "Dane CaseState (JSON):\n{case_state}\n\n"
                    "Na podstawie powyższych danych przygotuj ActionPlan (lista 'actions')."
                ),
            ),
        ]
    )

    chain = prompt | llm | parser

    try:
        result: ActionPlan = chain.invoke({
            "case_state": case_state.model_dump_json(),
            "format_instructions": parser.get_format_instructions(),
            "guide_excerpt": guide_excerpt,
        })
        return result.actions
    except Exception as e:
        print(f"Błąd generowania zaleceń: {e}")
        return []

PESEL_REGEX = re.compile(r"^\d{11}$")
POSTAL_REGEX = re.compile(r"\d{2}-\d{3}")
NIP_REGEX = re.compile(r"^\d{10}$")
REGON_REGEX = re.compile(r"^(\d{9}|\d{14})$")
PKD_REGEX = re.compile(r"^\d{2}\.\d{2}\.[A-Z]$", re.IGNORECASE)

def check_validation_of_fields(state: CaseState) -> list[str]:
    alerts: list[str] = []
    if state.pesel and not PESEL_REGEX.match(state.pesel):
        alerts.append("PESEL musi mieć dokładnie 11 cyfr.")
    if state.nip:
        normalized_nip = state.nip.replace(" ", "").replace("-", "")
        if not NIP_REGEX.match(normalized_nip):
            alerts.append("NIP powinien zawierać 10 cyfr (bez spacji i kresek).")
    if state.regon:
        normalized_regon = state.regon.replace(" ", "").replace("-", "")
        if not REGON_REGEX.match(normalized_regon):
            alerts.append("REGON powinien mieć 9 lub 14 cyfr.")
    if state.pkd:
        normalized_pkd = state.pkd.strip().upper()
        if not PKD_REGEX.match(normalized_pkd):
            alerts.append("PKD wpisz w formacie 00.00.X (np. 62.01.Z).")
    if state.address_home:
        match = POSTAL_REGEX.search(state.address_home)
        if not match:
            alerts.append("Adres zamieszkania powinien zawierać kod pocztowy w formacie 00-000.")
    return alerts

def prepend_validation_warnings(reply: str, alerts: list[str]) -> str:
    if not alerts:
        return reply
    warning = "Uwaga: wykryto problemy z formatem danych. Sprawdź proszę:\n- " + "\n- ".join(alerts)
    return f"{warning}\n\n{reply}" if reply else warning

def run_assistant_pipeline(
    case_id: str,
    message: str,
    mode: Mode,
    previous_state: Optional[CaseState] = None,
    conversation_history: Optional[List[ChatTurn]] = None,
) -> AssistantMessageResponse:
    """
    Miejsce na LangChain:
    - tutaj w przyszłości:
      - odczytasz CaseState z bazy na podstawie case_id,
      - uruchomisz chain do ekstrakcji informacji z message + historii rozmowy,
      - zaktualizujesz CaseState,
      - policzysz brakujące pola,
      - wygenerujesz kolejne pytania i draft dokumentu.

    Póki co używamy prostego chaina:
    - LLM uzupełnia CaseState na podstawie wiadomości,
    - prosta funkcja Pythonowa wykrywa braki.
    """
    # TODO: tutaj podłącz w przyszłości storage (np. bazę danych) po case_id
    base_state = previous_state or CaseState()

    # Dzisiejsza data z punktu widzenia backendu (ISO)
    today = date.today().isoformat()

    history = conversation_history or []

    # LangChain: próba uzupełnienia CaseState na podstawie wiadomości i historii
    case_state = extract_case_state_with_llm(
        previous_state=base_state,
        message=message,
        mode=mode,
        today=today,
        conversation_history=history,
    )

    case_state.address_home = normalize_address(case_state.address_home)
    case_state.address_correspondence = normalize_address(case_state.address_correspondence)
    validation_alerts = check_validation_of_fields(case_state)

    # Wyznacz pola, których użytkownik nie chce podawać – na podstawie historii + bieżącej odpowiedzi.
    skipped_from_history = set(infer_skipped_fields_from_history(history))
    # Sprawdź, czy bieżąca wiadomość jest odmową odpowiedzi lub prośbą o przejście
    # do kolejnej kategorii na podstawie ostatniego pytania asystenta.
    skipped_current: set[str] = set()
    if history:
        last_assistant = next(
            (t for t in reversed(history) if t.role == "assistant"), None
        )
        if last_assistant:
            text = last_assistant.content

            # Jeśli użytkownik prosi o przejście do następnej kategorii,
            # pomijamy wszystkie jeszcze niewypełnione pola z bieżącej kategorii.
            if message_asks_next_category(message):
                cat_label = extract_category_label_from_text(text)
                if cat_label:
                    cat = find_category_by_label(cat_label)
                    if cat:
                        _, _, category_fields = cat
                        for name in category_fields:
                            value = getattr(case_state, name, None)
                            if value is None or (
                                isinstance(value, str) and not value.strip()
                            ):
                                skipped_current.add(name)
            else:
                # Standardowy przypadek: sprawdzamy odmowę dla konkretnego pola
                if ":" in text:
                    label = (
                        text.split(":")[-2 if text.count(":") > 1 else 0]
                        .splitlines()[-1]
                        .strip()
                    )
                    field_name = field_name_from_label(label)
                    if field_name and detect_skip_with_llm(label, message):
                        skipped_current.add(field_name)

    skipped_all = list(skipped_from_history | skipped_current)

    # Sprawdź braki obowiązkowe (dla missing_fields), ignorując pola, które użytkownik świadomie pominął.
    missing = simple_missing_fields(case_state, mode, skipped_fields=skipped_all)

    # Kolejność WSZYSTKICH pól, o które chcemy po kolei pytać,
    # tak długo jak użytkownik nie odmówi (skipped_fields) i pole jest puste.
    question_order = [
        "reporter_type",
        # Dane poszkodowanego
        "first_name",
        "last_name",
        "pesel",
        "date_of_birth",
        "address_home",
        "address_correspondence",
        # Dane działalności
        "nip",
        "regon",
        "business_address",
        "pkd",
        "business_description",
        # Informacje o wypadku
        "accident_date",
        "accident_time",
        "accident_place",
        "planned_work_start",
        "planned_work_end",
        "injury_type",
        "accident_description",
        "first_aid_info",
        "proceedings_info",
        "equipment_info",
    ]

    skipped = set(skipped_all)
    next_field = None
    for name in question_order:
        if name in skipped:
            continue
        value = getattr(case_state, name, None)
        if value is None or (isinstance(value, str) and not value.strip()):
            next_field = name
            break

    if next_field:
        if next_field == "reporter_type":
            assistant_reply = (
                "Dzień dobry. Jestem wirtualnym asystentem ZUS.\n\n"
                "Zanim zaczniemy: Czy zgłaszasz wypadek osobiście (jako poszkodowany), "
                "czy występujesz jako pełnomocnik?"
            )
        elif next_field == "accident_description":
            assistant_reply = (
                "Przejdźmy do najważniejszej części - przebiegu wypadku. "
                "Aby ZUS mógł prawidłowo ocenić zdarzenie, musimy ustalić tzw. drzewo przyczyn.\n\n"
                "Opisz proszę sytuację, odpowiadając kolejno na trzy pytania:\n"
                "1. Co robiłeś/aś bezpośrednio przed zdarzeniem? (Jaka to była czynność? np. 'wchodziłem po drabinie', 'niosłem paczkę')\n"
                "2. Co dokładnie się stało? (Jaki fakt spowodował uraz? np. 'noga ześlizgnęła się ze stopnia', 'potknąłem się o przewód')\n"
                "3. Jaki jest skutek? (Jaki to uraz i jakiej części ciała?)\n\n"
                "Możesz to opisać własnymi słowami w jednej wiadomości."
            )
        else:

            category = find_category_for_field(next_field)
            if category is not None:
                _, category_label, category_fields = category
                # W tej kategorii pytamy tylko o pola, które nadal są puste
                # (i nie zostały oznaczone jako pominięte).
                missing_in_category = []
                for name in category_fields:
                    if name in skipped:
                        continue
                    value = getattr(case_state, name, None)
                    if value is None or (isinstance(value, str) and not value.strip()):
                        missing_in_category.append(name)

                labels = [human_field_label(name) for name in missing_in_category]
                fields_list = ", ".join(labels)

                # Pierwsza interakcja — przedstaw zasady raz na początku.
                if not history:
                    assistant_reply = (
                        "Dzień dobry! Opisz proszę, co się wydarzyło.\n\n"
                        "Potem przejdziemy po kilku grupach pytań potrzebnych do formularza. "
                        "Możesz mówić swobodnie, tak jak Ci wygodnie – jeśli czegoś nie chcesz podawać, "
                        "napisz po prostu, że tę informację pomijamy albo że wolisz następną kategorię.\n\n"
                        f"Na początek zatrzymajmy się przy kategorii: {category_label}. "
                        f"Napisz w kilku zdaniach, co uważasz za ważne w tym temacie. "
                        f"Jeśli chcesz, możesz przy okazji wspomnieć o rzeczach typu: {fields_list}."
                    )
                else:
                    # Kolejne pytania – jedna grupa pól na raz, w luźniejszej formie.
                    assistant_reply = (
                        f"Teraz kategoria: {category_label}. "
                        "Napisz po prostu, co uważasz za ważne w tym obszarze. "
                        f"Jeśli chcesz, możesz też doprecyzować rzeczy typu: {fields_list}. "
                        "Możesz pominąć dowolne elementy albo poprosić o następną kategorię."
                    )
            else:
                # Gdyby pole nie należało do żadnej kategorii (nie powinno się zdarzyć).
                label = human_field_label(next_field)
                assistant_reply = f"{label}:"
    else:
        # To jest moment, w którym generujemy raport końcowy i zalecenia
        
        # Wywołaj LLM z promptem:
        # "Na podstawie zgromadzonych danych (CaseState):
        # 1. Czy to był wypadek w drodze, czy przy pracy?
        # 2. Jakie dokumenty są wymagane dla tego konkretnego przypadku?
        # 3. Jakie kroki musi podjąć użytkownik?
        # Zwróć listę ActionStep."
        
        actions = generate_post_accident_actions(case_state) # Nowa funkcja z LLM
        
        assistant_reply = (
            "Dziękuję, to wszystkie pytania o dane wymagane w formularzu. "
            "Po pobraniu dokumentów wyślę ci listę kroków, które powinieneś teraz podjąć."
        )

        assistant_reply = prepend_validation_warnings(assistant_reply, validation_alerts)

        return AssistantMessageResponse(
            assistant_reply=assistant_reply,
            missing_fields=missing,
            case_state_preview=case_state,
            recommended_actions=actions
        )

    assistant_reply = prepend_validation_warnings(assistant_reply, validation_alerts)

    return AssistantMessageResponse(
        assistant_reply=assistant_reply,
        missing_fields=missing,
        case_state_preview=case_state,
    )


def evaluate_case_from_documents(
    documents: List[CaseDocument], case_id: str
) -> CaseEvaluationResult:
    """
    Uruchamia złożony pipeline LLM na komplecie dokumentów:
    - ujednolica stan faktyczny (CaseState),
    - wykrywa rozbieżności,
    - wskazuje brakujące informacje i dokumenty,
    - wydaje opinię, czy zdarzenie jest wypadkiem podczas prowadzenia pozarolniczej DG,
    - generuje projekt karty wypadku.
    """
    llm = get_llm()
    if llm is None or ChatPromptTemplate is None or PydanticOutputParser is None:
        # Fallback: minimalna implementacja bez LLM – zwracamy puste rozstrzygnięcie.
        base_state = CaseState()
        missing = simple_missing_fields(base_state, Mode.NOTIFICATION)
        return CaseEvaluationResult(
            normalized_case_state=base_state,
            discrepancies=[],
            missing_fields=missing,
            missing_documents=[],
            opinion=OpinionOutcome.INCONCLUSIVE,
            opinion_explanation=(
                "Brak dostępnego modelu LLM – nie można przeprowadzić pełnej oceny na podstawie dokumentów."
            ),
            accident_card_draft=(
                "Brak projektu karty wypadku – środowisko LLM nie jest dostępne."
            ),
        )

    parser = PydanticOutputParser(pydantic_object=CaseEvaluationResult)

    docs_as_text = [
        {
            "name": d.name,
            "type": d.type or "",
            "text": d.text,
        }
        for d in documents
    ]

    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                (
                    "Jesteś ekspertem ZUS ds. wypadków przy prowadzeniu pozarolniczej "
                    "działalności gospodarczej.\n"
                    "Na podstawie kompletu dokumentów musisz:\n"
                    "1) odczytać i zrozumieć treść dokumentów (zakładamy, że OCR już został wykonany),\n"
                    "2) ujednolicić stan faktyczny sprawy w strukturze CaseState,\n"
                    "3) wskazać rozbieżności pomiędzy dokumentami (np. różne daty, różne miejsca wypadku, "
                    "różne dane świadków, różne opisy zdarzenia),\n"
                    "4) wskazać brakujące informacje i/lub brakujące dokumenty, które są potrzebne do "
                    "rzetelnej oceny zdarzenia,\n"
                    "5) wydać jednoznaczną opinię, czy zdarzenie jest wypadkiem podczas prowadzenia "
                    "pozarolniczej działalności gospodarczej,\n"
                    "6) szczegółowo uzasadnić swoją opinię,\n"
                    "7) przygotować projekt karty wypadku (accident_card_draft) – może być jako dobrze "
                    "sformatowany tekst z nagłówkami i polami.\n\n"
                    "Zawsze zwracaj wynik w formacie JSON ściśle zgodnym ze schematem CaseEvaluationResult."
                ),
            ),
            (
                "human",
                (
                    "Identyfikator sprawy: {case_id}\n\n"
                    "Komplet dokumentów (po OCR) w formacie JSON:\n"
                    "{documents_json}\n\n"
                    "Twoje zadanie:\n"
                    "- przeanalizuj wszystkie dokumenty razem (załóż, że mogą zawierać błędy i sprzeczne dane),\n"
                    "- ujednolicz stan faktyczny w strukturze CaseState,\n"
                    "- wskaż wyraźnie wszystkie istotne rozbieżności pomiędzy dokumentami,\n"
                    "- wskaż, jakich informacji lub dokumentów brakuje,\n"
                    "- wydaj opinię (opinion) oraz szczegółowe uzasadnienie (opinion_explanation),\n"
                    "- przygotuj projekt karty wypadku (accident_card_draft) – może być jako dobrze "
                    "sformatowany tekst z nagłówkami i polami.\n"
                    "Odpowiedz TYLKO JSON-em zgodnym ze schematem CaseEvaluationResult."
                ),
            ),
        ]
    )

    chain = prompt | llm | parser

    try:
        result: CaseEvaluationResult = chain.invoke(
            {
                "case_id": case_id,
                "documents_json": docs_as_text,
            }
        )
        return result
    except Exception:
        # Bezpieczny fallback: nie przerywamy działania API, tylko zwracamy odpowiedź INCONCLUSIVE.
        base_state = CaseState()
        missing = simple_missing_fields(base_state, Mode.NOTIFICATION)
        return CaseEvaluationResult(
            normalized_case_state=base_state,
            discrepancies=[],
            missing_fields=missing,
            missing_documents=[],
            opinion=OpinionOutcome.INCONCLUSIVE,
            opinion_explanation=(
                "Wystąpił błąd podczas przetwarzania dokumentów przez model LLM – "
                "nie można przeprowadzić pełnej oceny."
            ),
            accident_card_draft=(
                "Projekt karty wypadku nie został wygenerowany z powodu błędu LLM."
            ),
        )


@app.post("/api/assistant/message", response_model=AssistantMessageResponse)
async def assistant_message(
    payload: AssistantMessageRequest,
) -> AssistantMessageResponse:
    """
    Główny endpoint czatu:
    - przyjmuje wiadomość użytkownika,
    - uruchamia pipeline asystenta,
    - zwraca tekst odpowiedzi, listę braków i podgląd stanu sprawy.
    """
    return run_assistant_pipeline(
        case_id=payload.case_id,
        message=payload.message,
        mode=payload.mode,
        previous_state=payload.case_state,
        conversation_history=payload.conversation_history,
    )


@app.post("/api/case/evaluate-documents", response_model=CaseEvaluationResponse)
async def evaluate_documents(payload: CaseEvaluationRequest) -> CaseEvaluationResponse:
    """
    Endpoint dla II etapu oceny:
    - przyjmuje komplet dokumentów w formie tekstowej (po OCR),
    - uruchamia LLM do analizy,
    - zwraca ujednolicony stan faktyczny, rozbieżności, braki, opinię i projekt karty wypadku.
    """
    evaluation = evaluate_case_from_documents(
        documents=payload.documents, case_id=payload.case_id
    )
    return CaseEvaluationResponse(case_id=payload.case_id, evaluation=evaluation)


@app.post("/api/ocr/read-document")
async def read_document_ocr(file: UploadFile = File(...)) -> dict:
    """
    OCR endpoint.

    Accepts an uploaded image file (e.g. PNG/JPEG) and returns text extracted
    from the document using Tesseract OCR.
    """
    contents = await file.read()
    if not contents:
        raise HTTPException(status_code=400, detail="Empty file")

    try:
        text = extract_text_from_image(contents)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:  # pragma: no cover - defensive
        raise HTTPException(status_code=500, detail="OCR failed") from exc

    return {
        "filename": file.filename,
        "text": text,
    }



# --- GENEROWANIE DOKUMENTÓW ---

class PDF(FPDF):
    def header(self):
        # Prosty nagłówek PDF
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'Dokument wygenerowany przez asystenta ZANT', 0, 1, 'C')
        self.ln(10)

def sanitize_text(text: Optional[str]) -> str:
    """Czyści tekst z None i obsługuje polskie znaki dla prostego PDF."""
    if text is None:
        return "---"
    return text

def create_notification_docx(state: CaseState) -> io.BytesIO:
    """Tworzy plik Word z Zawiadomieniem o wypadku."""
    doc = DocxDocument()
    
    doc.add_heading('ZAWIADOMIENIE O WYPADKU', 0)
    doc.add_heading('przy prowadzeniu działalności gospodarczej', 1)

    # Sekcja 1: Poszkodowany
    doc.add_heading('1. Dane poszkodowanego (Zgłaszającego)', level=2)
    p = doc.add_paragraph()
    p.add_run(f"Imię i nazwisko: ").bold = True
    p.add_run(f"{sanitize_text(state.first_name)} {sanitize_text(state.last_name)}\n")
    p.add_run(f"PESEL: ").bold = True
    p.add_run(f"{sanitize_text(state.pesel)}\n")
    p.add_run(f"Adres zamieszkania: ").bold = True
    p.add_run(f"{sanitize_text(state.address_home)}")

    # Sekcja 2: Płatnik składek
    doc.add_heading('2. Dane płatnika składek (Działalność)', level=2)
    p = doc.add_paragraph()
    p.add_run(f"Nazwa/Opis: ").bold = True
    p.add_run(f"{sanitize_text(state.business_description)}\n")
    p.add_run(f"NIP: ").bold = True
    p.add_run(f"{sanitize_text(state.nip)}   ")
    p.add_run(f"REGON: ").bold = True
    p.add_run(f"{sanitize_text(state.regon)}\n")
    p.add_run(f"Adres działalności: ").bold = True
    p.add_run(f"{sanitize_text(state.business_address)}")

    # Sekcja 3: Informacje o wypadku
    doc.add_heading('3. Informacje o wypadku', level=2)
    p = doc.add_paragraph()
    p.add_run(f"Data i godzina: ").bold = True
    p.add_run(f"{sanitize_text(state.accident_date)}, godz. {sanitize_text(state.accident_time)}\n")
    p.add_run(f"Miejsce wypadku: ").bold = True
    p.add_run(f"{sanitize_text(state.accident_place)}\n")
    p.add_run(f"Rodzaj urazu: ").bold = True
    p.add_run(f"{sanitize_text(state.injury_type)}")

    doc.add_heading('Okoliczności i przyczyny wypadku:', level=3)
    doc.add_paragraph(sanitize_text(state.accident_description))

    # Sekcja 4: Świadkowie
    if state.witnesses:
        doc.add_heading('4. Świadkowie', level=2)
        for i, w in enumerate(state.witnesses, 1):
            doc.add_paragraph(f"{i}. {w.first_name} {w.last_name}, adres: {w.address or 'brak'}")
    
    doc.add_paragraph("\n\n......................................................\n(podpis zgłaszającego)")

    # Zapis do bufora
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def create_explanation_docx(state: CaseState) -> io.BytesIO:
    """Tworzy plik Word z Wyjaśnieniami poszkodowanego."""
    doc = DocxDocument()
    doc.add_heading('ZAPIS WYJAŚNIEŃ POSZKODOWANEGO', 0)

    p = doc.add_paragraph()
    p.add_run(f"Ja, niżej podpisany/a: ").bold = True
    p.add_run(f"{sanitize_text(state.first_name)} {sanitize_text(state.last_name)}\n")
    
    doc.add_heading('Treść wyjaśnień:', level=2)
    # Tutaj wstawiamy opis wypadku jako główne wyjaśnienie
    doc.add_paragraph(sanitize_text(state.accident_description))

    if state.equipment_info:
         doc.add_heading('Informacje dodatkowe (maszyny, BHP):', level=3)
         doc.add_paragraph(state.equipment_info)

    doc.add_paragraph("\n\nOświadczam, że powyższe wyjaśnienia są zgodne z prawdą.")
    doc.add_paragraph("\n......................................................\n(data i podpis)")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def remove_polish_chars(text: str) -> str:
    """
    Zamienia polskie znaki i inne znaki specjalne na odpowiedniki ASCII/Latin-1,
    aby uniknąć błędów FPDF przy braku czcionki Unicode.
    """
    replacements = {
        'ą': 'a', 'ć': 'c', 'ę': 'e', 'ł': 'l', 'ń': 'n', 'ó': 'o', 'ś': 's', 'ź': 'z', 'ż': 'z',
        'Ą': 'A', 'Ć': 'C', 'Ę': 'E', 'Ł': 'L', 'Ń': 'N', 'Ó': 'O', 'Ś': 'S', 'Ź': 'Z', 'Ż': 'Z',
        '„': '"', '”': '"', '–': '-', '—': '-', '’': "'", '…': '...'
    }
    # 1. Znane zamienniki
    text = "".join(replacements.get(c, c) for c in text)
    
    # 2. Usuń wszystko co nie jest latin-1 (zamień na ?)
    # FPDF (standard fonts) obsługuje latin-1.
    return text.encode('latin-1', 'replace').decode('latin-1')


def create_simple_pdf(title: str, content_dict: dict) -> io.BytesIO:
    """
    Tworzy prosty PDF. Może zajmować wiele stron.
    Zastosowano bezpieczny layout wertykalny (Etykieta nad Wartością) dla długich pól,
    aby uniknąć błędów FPDF i wychodzenia poza margines.
    """
    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # Fallback dla standardowego Ariala
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, remove_polish_chars(title), ln=True, align='C')
    pdf.ln(5)

    pdf.set_font("Arial", size=11)
    
    # Efektywna szerokość strony (obszar roboczy)
    effective_page_width = pdf.w - pdf.l_margin - pdf.r_margin

    placeholder_line = ". " * 65

    for key, value in content_dict.items():
        clean_key = remove_polish_chars(str(key))
        
        # Nagłówki sekcji
        if key.startswith("---"):
            pdf.ln(5)
            pdf.set_font("Arial", 'B', 11)
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, 8, clean_key)
            pdf.set_font("Arial", '', 11)
            continue

        # Przygotowanie wartości
        extra_lines = 0
        placeholder_hint: Optional[str] = None
        field_value = value

        if isinstance(value, dict):
            extra_lines = int(value.get("lines") or 0)
            placeholder_hint = value.get("hint")
            field_value = value.get("value")

        if field_value and str(field_value).strip():
            clean_val = remove_polish_chars(str(field_value))
            if extra_lines:
                trailing_space = "\n".join(placeholder_line for _ in range(extra_lines))
                clean_val = f"{clean_val}\n{trailing_space}"
            is_placeholder = False
        else:
            # Placeholder z dodatkowymi liniami na wpisy odręczne
            total_lines = max(extra_lines, 2 if len(clean_key) > 25 else 1)
            block = "\n".join(placeholder_line for _ in range(total_lines))
            if placeholder_hint:
                clean_val = f"{remove_polish_chars(placeholder_hint)}\n{block}"
            else:
                clean_val = block
            is_placeholder = True

        # Layout
        pdf.set_font("Arial", 'B', 11)
        key_str = clean_key + ": "
        
        # Sprawdzamy szerokości
        key_width = pdf.get_string_width(key_str)
        
        # Ile miejsca zostanie na linii po wpisaniu klucza?
        remaining_space = effective_page_width - key_width
        
        # Logika decyzji o układzie:
        # 1. Placeholder -> nowa linia
        # 2. Wartość długa -> nowa linia
        # 3. Zostaje za mało miejsca na samej linii po kluczu (< 50mm) -> nowa linia (zapobiega crashom FPDF)
        
        force_new_line = (
            is_placeholder 
            or len(clean_val) > 60
            or remaining_space < 50
        )

        if force_new_line:
            # Wariant Wertykalny:
            # Etykieta
            # Wartość
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, 6, key_str)
            pdf.set_font("Arial", '', 11)
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, 6, clean_val)
        else:
            # Wariant Horyzontalny:
            # Etykieta: Wartość
            try:
                # Sprawdź czy nie jesteśmy na końcu strony
                if pdf.get_y() > pdf.h - pdf.b_margin - 10:
                    pdf.add_page()

                pdf.set_font("Arial", 'B', 11)
                pdf.set_x(pdf.l_margin)
                pdf.cell(key_width, 6, key_str, ln=False)
                
                pdf.set_font("Arial", '', 11)
                # multi_cell(0) używa pozostałej szerokości do prawego marginesu
                pdf.set_x(pdf.get_x())
                pdf.multi_cell(0, 6, clean_val)
            except Exception:
                # Fallback w razie błędu - resetujemy stan i drukujemy wertykalnie
                pdf.ln()
                pdf.set_x(pdf.l_margin)
                pdf.set_font("Arial", 'B', 11)
                pdf.multi_cell(0, 6, key_str)
                pdf.set_font("Arial", '', 11)
                pdf.multi_cell(0, 6, clean_val)

    output = io.BytesIO()
    output.write(pdf.output())
    output.seek(0)
    return output


def create_pdf_from_markdown(markdown_text: str, title: str = "Karta wypadku") -> io.BytesIO:
    """
    Bardzo prosta konwersja Markdown -> PDF oparta na istniejącym
    helperze `create_simple_pdf`.

    Nie próbuje odwzorować całego formatowania Markdown – traktuje
    wypełnioną kartę jako jeden długi blok tekstu w PDF.
    """
    content = {
        "Treść karty (Markdown)": {
            "value": markdown_text,
            # Kilka dodatkowych linii na ręczne dopiski
            "lines": 10,
        }
    }
    return create_simple_pdf(title, content)


# Dodatkowe wypełniajace EWYP
def parse_address_to_dict(address_str: Optional[str]) -> dict:
    """
    Rozbija string adresu (np. 'ul. Długa 5/10, 00-123 Warszawa') na części:
    ulica, nr_domu, nr_lokalu, kod, poczta.
    Działa na zasadzie heurystyki (best-effort).
    """
    result = {
        "ulica": "",
        "nr_domu": "",
        "nr_lokalu": "",
        "kod": "",
        "poczta": ""
    }
    
    if not address_str:
        return result

    # 1. Wyciągamy kod pocztowy (XX-XXX)
    zip_match = re.search(r'(\d{2}-\d{3})', address_str)
    if zip_match:
        result["kod"] = zip_match.group(1)
        # Poczta to zazwyczaj to, co jest po kodzie pocztowym
        parts = address_str.split(result["kod"])
        if len(parts) > 1:
            result["poczta"] = parts[1].strip(" ,.")
        
        # Reszta (przed kodem) to ulica i numer
        street_part = parts[0].strip(" ,.")
    else:
        street_part = address_str

    # 2. Próbujemy oddzielić numer domu i lokalu od ulicy
    # Szukamy cyfry na końcu ciągu
    # Wzorce: "Długa 5", "Długa 5/10", "Długa 5 m. 10"
    match_nr = re.search(r'(\d+[a-zA-Z]?\/?\s?\d*)$', street_part)
    if match_nr:
        numer_full = match_nr.group(1)
        # Ulica to wszystko przed numerem
        result["ulica"] = street_part[:match_nr.start()].strip()
        
        # Rozbicie na dom i lokal
        if '/' in numer_full:
            d, l = numer_full.split('/', 1)
            result["nr_domu"] = d.strip()
            result["nr_lokalu"] = l.strip()
        elif 'm.' in numer_full:
            d, l = numer_full.split('m.', 1)
            result["nr_domu"] = d.strip()
            result["nr_lokalu"] = l.strip()
        else:
            result["nr_domu"] = numer_full.strip()
    else:
        # Jeśli nie udało się znaleźć numeru, wszystko trafia do ulicy
        result["ulica"] = street_part

    return result


def normalize_address(address: Optional[str]) -> Optional[str]:
    """Formatuje adres do postaci 'Ulica 1/2, 00-000 Miasto'"""
    if not address or not address.strip():
        return address

    parts = parse_address_to_dict(address)
    if not any(parts.values()):
        return address.strip()

    street = parts["ulica"].strip()
    house = parts["nr_domu"].strip()
    flat = parts["nr_lokalu"].strip()
    postal = parts["kod"].strip()
    city = parts["poczta"].strip()

    street_block = street
    if house:
        number = house
        if flat:
            number = f"{number}/{flat}"
        street_block = f"{street} {number}".strip()

    locality = " ".join(item for item in [postal, city] if item).strip()

    pieces = [p for p in [street_block, locality] if p]
    normalized = ", ".join(pieces)
    return normalized or address.strip()


def fill_ewyp_pdf(case_state: CaseState, template_path: str = "EWYP.pdf") -> io.BytesIO:
    """
    Wypełnia formularz ZUS EWYP (wersja naprawiona: krótkie nazwy pól).
    Rozróżnia strony, aby unikać konfliktów (np. Imię[0] na str 1 vs str 5).
    """
    reader = PdfReader(template_path)
    writer = PdfWriter()
    writer.append_pages_from_reader(reader)
    
    # Funkcje pomocnicze
    def val(v): return str(v) if v is not None else ""

    def date_val(iso_date: Optional[str]) -> str:
        if not iso_date:
            return ""
        # Jeśli data jest w formacie YYYY-MM-DD (np. 2000-12-07)
        parts = str(iso_date).split('-')
        if len(parts) == 3 and len(parts[0]) == 4:
            # Zwracamy DD-MM-YYYY (np. 07-12-2000)
            return f"{parts[2]}{parts[1]}{parts[0]}"
        return str(iso_date)
    
    # Parsowanie adresów
    ah = parse_address_to_dict(case_state.address_home)   # Adres domowy (Page 1)
    ac = parse_address_to_dict(case_state.address_correspondence) # Adres korespondencji (Page 2)
    
    # --- LOGIKA CHECKBOXÓW ---
    def get_chk_pair(condition):
        # Zwraca krotkę (wartość_dla_TAK, wartość_dla_NIE)
        # W Twoim logu nie widać wartości eksportu (On/Yes), ale standardowo to /Yes i /Off
        if condition: return ('/Yes', '/Off')
        return ('/Off', '/Yes')

    has_aid = case_state.first_aid_info is not None and len(str(case_state.first_aid_info)) > 2
    chk_aid_tak, chk_aid_nie = get_chk_pair(has_aid)

    has_eq = case_state.equipment_info is not None and len(str(case_state.equipment_info)) > 2
    chk_eq_tak, chk_eq_nie = get_chk_pair(has_eq)

    # --- PĘTLA PO STRONACH ---
    # Definiujemy mapę pól ODDZIELNIE dla każdej strony (0-indexed),
    # aby Imię[0] na str. 1 (poszkodowany) nie nadpisało Imię[0] na str. 5 (świadek).
    
    for page_num, page in enumerate(writer.pages):
        # Słownik z danymi tylko dla TEJ konkretnej strony
        page_map = {}

        if page_num == 0:  # --- STRONA 1: DANE POSZKODOWANEGO ---
            page_map = {
                'PESEL[0]': val(case_state.pesel),
                'Imię[0]': val(case_state.first_name),
                'Nazwisko[0]': val(case_state.last_name),
                'Dataurodzenia[0]': date_val(case_state.date_of_birth),
                # Adres zamieszkania
                'Ulica[0]': ah['ulica'],
                'Numerdomu[0]': ah['nr_domu'],
                'Numerlokalu[0]': ah['nr_lokalu'],
                'Kodpocztowy[0]': ah['kod'],
                'Poczta[0]': ah['poczta'],
            }

        elif page_num == 1: # --- STRONA 2: ADRES DZIAŁALNOŚCI ---
            # Tutaj mapujemy dane firmy na pola Ulica2, itd.
            page_map = {
                'Ulica2[1]': ac['ulica'],
                'Numerdomu2[1]': ac['nr_domu'],
                'Numerlokalu2[1]': ac['nr_lokalu'],
                'Kodpocztowy2[1]': ac['kod'],
                'Poczta2[1]': ac['poczta'],
            }

        elif page_num == 2: # --- STRONA 3: CZAS I MIEJSCE WYPADKU ---
            page_map = {
                'Datawyp[0]': date_val(case_state.accident_date),
                'Godzina[0]': val(case_state.accident_time),
                'Miejscewyp[0]': val(case_state.accident_place),
                'Godzina3A[0]': val(case_state.planned_work_start), # Rozpoczęcie
                'Godzina3B[0]': val(case_state.planned_work_end),   # Zakończenie
            }

        elif page_num == 3: # --- STRONA 4: OPISY I CHECKBOXY ---
            page_map = {
                'Tekst4[0]': val(case_state.injury_type),           # Urazy
                'Tekst5[0]': val(case_state.accident_description),  # Opis
                'Tekst6[0]': val(case_state.first_aid_info),        # Info o pomocy
                'Tekst7[0]': val(case_state.proceedings_info),      # Postępowanie
                'Tekst8[0]': val(case_state.equipment_info),        # Maszyny
                
                # Checkboxy (Pierwsza pomoc)
                'TAK6[0]': chk_aid_tak,
                'NIE6[0]': chk_aid_nie,
                
                # Checkboxy (Maszyny)
                'TAK8[0]': chk_eq_tak,
                'NIE8[0]': chk_eq_nie,
            }

        elif page_num == 4: # --- STRONA 5: ŚWIADKOWIE ---
            if case_state.witnesses:
                # Świadek 1
                w1 = case_state.witnesses[0]
                aw1 = parse_address_to_dict(w1.address)
                page_map.update({
                    'Imię[0]': val(w1.first_name),
                    'Nazwisko[0]': val(w1.last_name),
                    'Ulica[0]': aw1['ulica'],
                    'Numerdomu[0]': aw1['nr_domu'],
                    'Numerlokalu[0]': aw1['nr_lokalu'],
                    'Kodpocztowy[0]': aw1['kod'],
                    'Poczta[0]': aw1['poczta'],
                })
                
                # Świadek 2 (jeśli jest)
                if len(case_state.witnesses) > 1:
                    w2 = case_state.witnesses[1]
                    aw2 = parse_address_to_dict(w2.address)
                    # Wg loga na stronie 5 jest Imię[1], Nazwisko[1] oraz Ulica[1]
                    page_map.update({
                        'Imię[1]': val(w2.first_name),
                        'Nazwisko[1]': val(w2.last_name),
                        'Ulica[1]': aw2['ulica'],
                        'Numerdomu[1]': aw2['nr_domu'],
                        'Numerlokalu[1]': aw2['nr_lokalu'],
                        'Kodpocztowy[1]': aw2['kod'],
                        'Poczta[1]': aw2['poczta'],
                    })

        elif page_num == 5: # --- STRONA 6: DATA I PODPIS ---
             page_map = {
                 'Data[0]': date_val(case_state.accident_date) # lub date.today().isoformat()
             }

        # --- APLIKOWANIE DANYCH (LOW LEVEL) ---
        if "/Annots" in page:
            for annot in page["/Annots"]:
                try:
                    obj = annot.get_object()
                    if "/T" in obj:
                        key = obj["/T"]
                        
                        if key in page_map:
                            val_to_set = page_map[key]
                            
                            # Logika dla Checkboxów vs Tekst
                            if val_to_set in ['/Yes', '/Off']:
                                obj[NameObject("/V")] = NameObject(val_to_set)
                                obj[NameObject("/AS")] = NameObject(val_to_set)
                            else:
                                obj[NameObject("/V")] = TextStringObject(val_to_set)
                            
                            # KLUCZOWE: Usuwamy /AP, aby wymusić renderowanie tekstu
                            if "/AP" in obj:
                                del obj["/AP"]
                except Exception:
                    continue # Ignorujemy błędy pojedynczych pól

    # --- FIX STRUKTURY FORMULARZA (GLOBALNY) ---
    if "/AcroForm" not in writer.root_object:
        writer.root_object[NameObject("/AcroForm")] = DictionaryObject()
    
    acroform = writer.root_object["/AcroForm"]
    
    # Usuwamy XFA (dynamiczny formularz), zostawiamy tylko nasze dane statyczne
    if "/XFA" in acroform:
        del acroform["/XFA"]
    if "/XFA" in writer.root_object: # Czasem jest też w root
        del writer.root_object["/XFA"]

    # Wymuszamy na Adobe Readerze przeliczenie wyglądu
    acroform[NameObject("/NeedAppearances")] = BooleanObject(True)

    out = io.BytesIO()
    writer.write(out)
    out.seek(0)
    return out


@app.post("/api/case/download-documents")
async def download_documents(case_state: CaseState):
    """
    Generuje komplet dokumentów:
    1. ZUS EWYP (PDF) - wypełniony danymi formularz
    2. Wyjaśnienia (DOCX) - edytowalne
    3. Wyjaśnienia (PDF) - załącznik
    4. Wszystkie Dane z formularza
    i zwraca je jako ZIP.
    """
    
    files_to_zip = []
    actions = generate_post_accident_actions(case_state)

    # 1. Zawiadomienie o wypadku (Oryginalny PDF ZUS)
    try:
        # Zakładamy, że plik EWYP.pdf leży w głównym katalogu
        pdf_zus_content = fill_ewyp_pdf(case_state, template_path="EWYP.pdf")
        files_to_zip.append(("ZUS_EWYP_Zgloszenie.pdf", pdf_zus_content))
    except Exception as e:
        print(f"Błąd wypełniania PDF ZUS: {e}")
        # W razie błędu nie przerywamy, po prostu nie dodajemy tego pliku
        # lub dodajemy pusty plik error.txt
        pass

    # 2. Wyjaśnienia (DOCX) - zawsze warto mieć edytowalne
    if case_state.accident_description:
        docx_explanation = create_explanation_docx(case_state)
        files_to_zip.append(("Wyjasnienia_Poszkodowanego.docx", docx_explanation))
        
        # 3. Wyjaśnienia (PDF) - jako prosty załącznik
        explanation_data = {
            "--- WYJAŚNIENIA POSZKODOWANEGO ---": "",
            "Imię i nazwisko": f"{case_state.first_name or ''} {case_state.last_name or ''}".strip() or None,
            "Treść wyjaśnień": case_state.accident_description,
            "Informacje dodatkowe": case_state.equipment_info,
        }
        # Używamy Twojej funkcji create_simple_pdf (upewnij się, że jest w kodzie)
        pdf_explanation = create_simple_pdf("Załącznik - Wyjaśnienia", explanation_data)
        files_to_zip.append(("Zalacznik_Wyjasnienia.pdf", pdf_explanation))

    reporter_label = None
    if case_state.reporter_type == ReporterType.VICTIM:
        reporter_label = "Poszkodowany"
    elif case_state.reporter_type == ReporterType.PROXY:
        reporter_label = "Pełnomocnik"

    proxy_info = None
    if case_state.proxy_document_attached is True:
        proxy_info = "Tak"
    elif case_state.proxy_document_attached is False:
        proxy_info = "Nie"

    witness_hint = "Imię Nazwisko, Adres (oddzieleni przecinkami)"
    witness_value = None
    if case_state.witnesses:
        witness_rows = []
        for w in case_state.witnesses:
            name = f"{w.first_name or ''} {w.last_name or ''}".strip()
            entry = ", ".join(
                part for part in [name, w.address or None] if part and part.strip()
            )
            if entry:
                witness_rows.append(entry)
        witness_value = "\n".join(witness_rows) or None

    notification_data = {
        "--- DANE ZGŁASZAJĄCEGO ---": "",
        "Kto zgłasza wypadek": reporter_label,
        "Czy załączono pełnomocnictwo": proxy_info,

        "--- DANE POSZKODOWANEGO ---": "",
        "Imię": case_state.first_name,
        "Nazwisko": case_state.last_name,
        "PESEL": case_state.pesel,
        "Data urodzenia": case_state.date_of_birth,
        "Adres zamieszkania": case_state.address_home,
        "Adres do korespondencji": case_state.address_correspondence,

        "--- DANE PŁATNIKA SKŁADEK ---": "",
        "NIP": case_state.nip,
        "REGON": case_state.regon,
        "Adres siedziby": case_state.business_address,
        "Kod PKD": case_state.pkd,
        "Opis działalności": {"value": case_state.business_description, "lines": 2},

        "--- INFORMACJE O WYPADKU ---": "",
        "Data wypadku": case_state.accident_date,
        "Godzina wypadku": case_state.accident_time,
        "Miejsce wypadku": {"value": case_state.accident_place, "lines": 2},
        "Rodzaj urazu": case_state.injury_type,
        "Opis zdarzenia": {"value": case_state.accident_description, "lines": 5},
        "Udzielona pierwsza pomoc": {"value": case_state.first_aid_info, "lines": 3},
        "Postępowanie powypadkowe": {"value": case_state.proceedings_info, "lines": 3},
        "Maszyny i urządzenia": {"value": case_state.equipment_info, "lines": 3},

        "--- CZAS PRACY ---": "",
        "Planowane rozpoczęcie pracy": case_state.planned_work_start,
        "Planowane zakończenie pracy": case_state.planned_work_end,

        "--- ŚWIADKOWIE ---": "",
        "Dane świadków": {"value": witness_value, "lines": 4, "hint": witness_hint},
    }

    pdf_notification = create_simple_pdf("Zawiadomienie (Draft PDF)", notification_data)
    files_to_zip.append(("zawiadomienie_o_wypadku.pdf", pdf_notification))

    # --- Pakowanie do ZIP ---
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for filename, file_stream in files_to_zip:
            file_stream.seek(0)
            zip_file.writestr(filename, file_stream.read())

    zip_buffer.seek(0)

    # Nazwa pliku wynikowego
    filename = f"dokumenty_wypadkowe_{case_state.last_name or 'draft'}.zip"
    
    headers = {"Content-Disposition": f"attachment; filename={filename}"}

    serialized = json.dumps([action.model_dump() for action in actions], ensure_ascii=False)
    headers["X-ZANT-Actions"] = base64.b64encode(serialized.encode("utf-8")).decode("ascii")
    
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers=headers,
    )

@app.post("/api/ocr/summarize-accident-facts")
async def summarize_accident_facts(files: List[UploadFile] = File(...)) -> dict:
    """
    Przyjmuje wiele plików PDF z kartami wypadku, wykonuje OCR,
    a następnie zwraca zsyntetyzowane podsumowanie faktów istotnych
    dla oceny wypadku.
    """
    if not files:
        raise HTTPException(status_code=400, detail="No files uploaded")

    contents_list: list[bytes] = []
    for f in files:
        data = await f.read()
        if not data:
            continue
        contents_list.append(data)

    if not contents_list:
        raise HTTPException(status_code=400, detail="All uploaded files are empty")

    try:
        summary = summarize_accident_facts_from_pdfs(contents_list)
        filled_card_text = build_filled_card_text_from_summary(summary)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:  # pragma: no cover - defensive
        raise HTTPException(status_code=500, detail="Summarization failed") from exc

    pdf_base64: Optional[str] = None
    try:
        pdf_buffer = create_pdf_from_markdown(
            filled_card_text,
            title="Karta wypadku (draft)",
        )
        pdf_base64 = base64.b64encode(pdf_buffer.getvalue()).decode("ascii")
    except Exception:
        # Jeśli PDF z jakiegoś powodu się nie wygeneruje, nie blokujemy całej odpowiedzi.
        pdf_base64 = None

    return {
        "summary": summary,
        "file_count": len(contents_list),
        "accident_card_filled_text": filled_card_text,
        "accident_card_pdf_base64": pdf_base64,
    }
